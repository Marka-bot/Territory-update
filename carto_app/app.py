import base64
import json
from datetime import datetime
from io import BytesIO
from pathlib import Path

from flask import Flask, jsonify, render_template, request
import folium
from folium.plugins import Draw
import pdfkit
import qrcode
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches

BASE_DIR = Path(__file__).resolve().parent
EXPORT_PDF_DIR = BASE_DIR / "exports" / "cartes_pdf"
EXPORT_DOCX_DIR = BASE_DIR / "exports" / "cartes_docx"

for directory in (EXPORT_PDF_DIR, EXPORT_DOCX_DIR):
    directory.mkdir(parents=True, exist_ok=True)

app = Flask(__name__, template_folder=str(BASE_DIR / "templates"), static_folder=str(BASE_DIR / "static"))


def _safe_float(value, fallback):
    try:
        return float(value)
    except (TypeError, ValueError):
        return fallback


def generate_qr_base64(latitude: float, longitude: float) -> str:
    qr = qrcode.QRCode(version=1, box_size=10, border=2)
    qr.add_data(f"{latitude},{longitude}")
    qr.make(fit=True)
    image = qr.make_image(fill_color="black", back_color="white")
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return base64.b64encode(buffer.getvalue()).decode("utf-8")


def create_map(latitude: float, longitude: float):
    folium_map = folium.Map(location=[latitude, longitude], zoom_start=15, control_scale=True)
    folium.Marker([latitude, longitude], tooltip="Centre sélectionné").add_to(folium_map)

    draw = Draw(
        export=False,
        position="topleft",
        draw_options={
            "polyline": {"shapeOptions": {"color": "#ff0000", "weight": 3}},
            "polygon": {"shapeOptions": {"color": "#ff0000", "fillColor": "#ff0000"}},
            "rectangle": {"shapeOptions": {"color": "#ff0000", "fillColor": "#ff0000"}},
            "circle": False,
            "circlemarker": False,
            "marker": False,
        },
        edit_options={"edit": True, "remove": True},
    )
    draw.add_to(folium_map)

    map_id = folium_map.get_name()
    map_html = folium_map._repr_html_()
    return map_html, map_id


@app.route("/", methods=["GET", "POST"])
def index():
    default_latitude = 48.8566
    default_longitude = 2.3522

    latitude_input = request.form.get("latitude") if request.method == "POST" else request.args.get("latitude")
    longitude_input = request.form.get("longitude") if request.method == "POST" else request.args.get("longitude")

    latitude = _safe_float(latitude_input, default_latitude)
    longitude = _safe_float(longitude_input, default_longitude)

    map_html, map_id = create_map(latitude, longitude)
    qr_code_base64 = generate_qr_base64(latitude, longitude)

    return render_template(
        "index.html",
        latitude=latitude,
        longitude=longitude,
        map_html=map_html,
        map_id=map_id,
        qr_code=qr_code_base64,
    )


def _decode_base64_image(data_uri: str) -> bytes:
    if not data_uri:
        raise ValueError("Image data is required")
    if data_uri.startswith("data:image"):
        header, encoded = data_uri.split(",", 1)
    else:
        encoded = data_uri
    return base64.b64decode(encoded)


def _build_annotations_text(annotations: dict) -> str:
    if not annotations:
        return "Aucune annotation fournie."
    return json.dumps(annotations, indent=2, ensure_ascii=False)


def _create_pdf_with_pdfkit(pdf_path: Path, map_b64: str, qr_b64: str, latitude: float, longitude: float, annotations_text: str) -> bool:
    html_content = f"""
    <html>
        <head>
            <meta charset='utf-8'>
            <style>
                body {{ font-family: Arial, sans-serif; padding: 20px; }}
                .section {{ margin-bottom: 20px; }}
                .images {{ display: flex; gap: 20px; align-items: flex-start; }}
                img {{ max-width: 100%; height: auto; border: 1px solid #ccc; }}
                pre {{ background: #f8f8f8; padding: 10px; border-radius: 4px; }}
            </style>
        </head>
        <body>
            <h1>Carte annotée</h1>
            <div class="section">
                <strong>Coordonnées :</strong> {latitude}, {longitude}<br/>
                <strong>Date :</strong> {datetime.now().strftime("%d/%m/%Y %H:%M")}
            </div>
            <div class="section images">
                <div>
                    <h2>Carte</h2>
                    <img src="data:image/png;base64,{map_b64}" alt="Carte annotée" />
                </div>
                <div>
                    <h2>QR Code</h2>
                    <img src="data:image/png;base64,{qr_b64}" alt="QR Code" />
                </div>
            </div>
            <div class="section">
                <h2>Annotations</h2>
                <pre>{annotations_text}</pre>
            </div>
        </body>
    </html>
    """
    try:
        pdfkit.from_string(html_content, str(pdf_path))
        return True
    except (OSError, IOError):
        return False


def _create_pdf_with_pillow(pdf_path: Path, map_bytes: bytes, qr_bytes: bytes, latitude: float, longitude: float, annotations_text: str) -> None:
    map_image = Image.open(BytesIO(map_bytes)).convert("RGB")
    qr_image = Image.open(BytesIO(qr_bytes)).convert("RGB")

    max_qr_height = map_image.height // 3
    qr_ratio = qr_image.width / qr_image.height
    qr_new_height = min(max_qr_height, 300)
    qr_new_width = int(qr_new_height * qr_ratio)
    qr_resized = qr_image.resize((qr_new_width, qr_new_height))

    padding = 40
    text_area_height = 200
    width = max(map_image.width + padding * 2, qr_resized.width + padding * 2)
    height = map_image.height + qr_resized.height + text_area_height + padding * 4

    canvas = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(canvas)

    map_x = (width - map_image.width) // 2
    canvas.paste(map_image, (map_x, padding))

    qr_x = (width - qr_resized.width) // 2
    qr_y = map_image.height + padding * 2
    canvas.paste(qr_resized, (qr_x, qr_y))

    text_y = qr_y + qr_resized.height + padding
    text_content = f"Coordonnées : {latitude}, {longitude}\nDate : {datetime.now().strftime('%d/%m/%Y %H:%M')}\n\nAnnotations :\n{annotations_text}"
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 16)
    except OSError:
        font = ImageFont.load_default()

    draw.multiline_text((padding, text_y), text_content, fill="black", font=font, spacing=4)
    canvas.save(str(pdf_path), "PDF", resolution=100.0)


@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    data = request.get_json(force=True)
    image_data = data.get("imageData")
    annotations = data.get("annotations")
    latitude = _safe_float(data.get("latitude"), 0.0)
    longitude = _safe_float(data.get("longitude"), 0.0)

    map_bytes = _decode_base64_image(image_data)
    qr_base64 = generate_qr_base64(latitude, longitude)
    qr_bytes = base64.b64decode(qr_base64)

    annotations_text = _build_annotations_text(annotations)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    pdf_path = EXPORT_PDF_DIR / f"carte_{timestamp}.pdf"

    success = _create_pdf_with_pdfkit(pdf_path, base64.b64encode(map_bytes).decode("utf-8"), qr_base64, latitude, longitude, annotations_text)
    if not success:
        _create_pdf_with_pillow(pdf_path, map_bytes, qr_bytes, latitude, longitude, annotations_text)

    return jsonify({"success": True, "filename": pdf_path.name})


@app.route("/export/docx", methods=["POST"])
def export_docx():
    data = request.get_json(force=True)
    image_data = data.get("imageData")
    annotations = data.get("annotations")
    latitude = _safe_float(data.get("latitude"), 0.0)
    longitude = _safe_float(data.get("longitude"), 0.0)

    map_bytes = _decode_base64_image(image_data)
    qr_base64 = generate_qr_base64(latitude, longitude)
    qr_bytes = base64.b64decode(qr_base64)
    annotations_text = _build_annotations_text(annotations)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = EXPORT_DOCX_DIR / f"carte_{timestamp}.docx"

    document = Document()
    document.add_heading("Carte annotée", level=1)
    document.add_paragraph(f"Coordonnées : {latitude}, {longitude}")
    document.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    map_image_stream = BytesIO(map_bytes)
    document.add_picture(map_image_stream, width=Inches(5.5))

    document.add_paragraph("QR Code :")
    qr_stream = BytesIO(qr_bytes)
    document.add_picture(qr_stream, width=Inches(2))

    document.add_paragraph("Annotations :")
    document.add_paragraph(annotations_text)

    document.save(str(docx_path))

    return jsonify({"success": True, "filename": docx_path.name})


if __name__ == "__main__":
    app.run(debug=True)
